from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT_PATH = r"C:\Users\cees\Codebase\afstort\Gebruikershandleiding_Afstortportaal.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def style_run(run, size=11, bold=False, color="000000", font="Arial"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, after=6)
    run = p.add_run(text)
    style_run(run, size=11)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    set_paragraph_spacing(p, before=18 if level == 1 else 12, after=6)
    run = p.add_run(text)
    style_run(run, size={1: 18, 2: 14, 3: 12}[level], bold=False, color="000000")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, after=4)
    run = p.add_run(text)
    style_run(run, size=11)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(p, after=4)
    run = p.add_run(text)
    style_run(run, size=11)
    return p


def add_note_box(doc, title, lines):
    table = doc.add_table(rows=1 + len(lines), cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16)

    header = table.cell(0, 0)
    header.text = ""
    p = header.paragraphs[0]
    set_paragraph_spacing(p, before=0, after=2)
    run = p.add_run(title)
    style_run(run, size=11, bold=True)
    set_cell_shading(header, "F3E6D8")

    for idx, line in enumerate(lines, start=1):
        cell = table.cell(idx, 0)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, before=0, after=2)
        run = p.add_run(line)
        style_run(run, size=10)
        set_cell_shading(cell, "FBF7F1")

    doc.add_paragraph()


def add_screenshot_table(doc):
    rows = [
        ("1", "Login-scherm", "Inloggen", "E-mail, wachtwoord en inlogknop."),
        ("2", "2FA instellen", "2FA instellen", "QR-code, setup key en controlecodeveld."),
        ("3", "2FA controle", "Inloggen met 2FA", "Scherm voor code uit app of e-mail."),
        ("4", "Ritten-overzicht desktop", "Rittenlijst gebruiken", "Ritten met kleuren en hoofdvelden."),
        ("5", "Rit bevestigen", "Rit bevestigen", "Knop 'Bevestig rit' en liefst het bevestigingsvenster."),
        ("6", "Mobiele weergave", "Mobiele weergave", "Ritkaart met status, route- en contactknoppen."),
        ("7", "Afhaalbevestiging", "Afhaalbewijs en busbriefje", "Printbaar afhaalbewijs of document uit de e-mail."),
        ("8", "Rapport", "Rapport bekijken", "Rapport met afgeronde ritten en kilometers."),
        ("9", "E-mailscherm afronding", "Rit afronden", "Afrondingsmail en eventuele bijlagen."),
    ]

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    widths = [Cm(1.0), Cm(4.1), Cm(4.0), Cm(7.3)]

    hdr = table.rows[0].cells
    headers = ["Nr.", "Screenshot", "Plaats in handleiding", "Wat moet zichtbaar zijn"]
    for idx, text in enumerate(headers):
        hdr[idx].width = widths[idx]
        hdr[idx].text = ""
        p = hdr[idx].paragraphs[0]
        set_paragraph_spacing(p, before=0, after=0)
        run = p.add_run(text)
        style_run(run, size=9, bold=True)
        set_cell_shading(hdr[idx], "EFE7DC")

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].width = widths[idx]
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            set_paragraph_spacing(p, before=0, after=0)
            run = p.add_run(text)
            style_run(run, size=9)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)

styles = doc.styles
styles["Normal"].font.name = "Arial"
styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
styles["Normal"].font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_paragraph_spacing(title, before=0, after=3)
run = title.add_run("Gebruikershandleiding Afstortportaal")
style_run(run, size=24, bold=False)

subtitle = doc.add_paragraph()
set_paragraph_spacing(subtitle, before=0, after=10)
run = subtitle.add_run("Voor chauffeurs en collega's die willen begrijpen hoe de applicatie werkt")
style_run(run, size=11, color="555555")

intro = [
    "Deze handleiding helpt nieuwe gebruikers om met het afstortportaal te werken.",
    "De nadruk ligt op de dagelijkse werkwijze van chauffeurs.",
    "Daarnaast staat kort beschreven wat collega's binnen de organisatie in de applicatie kunnen doen en zien.",
]
for line in intro:
    add_body_paragraph(doc, line)

add_note_box(
    doc,
    "Kort samengevat",
    [
        "Chauffeurs gebruiken het portaal om ritten te bekijken, afspraken vast te leggen, bevestigingen te versturen en ritten af te ronden.",
        "Op telefoon is er een compacte mobiele weergave om ritten snel te bekijken en contact op te nemen.",
        "Voor het volledige beheren en afronden van ritten is de uitgebreide rittenlijst nodig.",
    ],
)

add_heading(doc, "1. Wat de applicatie doet", 1)
for text in [
    "Het afstortportaal ondersteunt het ophalen en afhandelen van collecte-opbrengsten.",
    "In de applicatie staat per rit wie de contactpersoon is, waar de opbrengst moet worden opgehaald, om welk soort geld het gaat en wat de voortgang van de rit is.",
    "De applicatie ondersteunt ook inloggen met tweestapsverificatie, het versturen van bevestigingsmails, het maken van afhaalbewijzen en het tonen van rapporten met afgeronde ritten.",
]:
    add_body_paragraph(doc, text)

add_heading(doc, "2. Belangrijke schermen", 1)
for text in [
    "De applicatie bestaat grofweg uit vier onderdelen: het login-scherm, de 2FA-schermen, de uitgebreide rittenlijst en de mobiele weergave 'Mijn ritten'.",
    "Voor chauffeurs is vooral het onderscheid tussen de uitgebreide rittenlijst en de mobiele weergave belangrijk.",
]:
    add_body_paragraph(doc, text)

add_bullet(doc, "De uitgebreide rittenlijst is bedoeld voor het plannen, bevestigen en afronden van ritten.")
add_bullet(doc, "De mobiele weergave is bedoeld om onderweg snel ritinformatie, route, telefoonnummer en e-mailadres te bekijken.")
add_bullet(doc, "Het rapport toont afgeronde ritten, gereden kilometers en het declarabele bedrag.")

add_heading(doc, "3. Inloggen", 1)
add_heading(doc, "3.1 Eerste stap: e-mail en wachtwoord", 2)
for text in [
    "Ga naar het afstortportaal en vul je e-mailadres en wachtwoord in.",
    "Na een juiste combinatie ga je niet direct naar de rittenlijst. Eerst volgt een extra controle via 2FA.",
]:
    add_body_paragraph(doc, text)
add_body_paragraph(doc, "Screenshot toevoegen: het login-scherm met e-mailveld, wachtwoordveld en de inlogknop.")

add_heading(doc, "3.2 2FA instellen", 2)
for text in [
    "Bij het eerste gebruik kan de applicatie vragen om 2FA in te stellen.",
    "Je kunt hiervoor een authenticator-app gebruiken, zoals Google Authenticator, Microsoft Authenticator, 1Password of Bitwarden.",
    "Op het scherm zie je een QR-code en een setup key. Scan de QR-code met je app en vul daarna de 6-cijferige controlecode in.",
]:
    add_body_paragraph(doc, text)
add_body_paragraph(doc, "Gebruik je liever geen authenticator-app, dan kun je de controlecode ook per e-mail ontvangen.")
add_body_paragraph(doc, "Screenshot toevoegen: het 2FA-instelvenster met QR-code, setup key en controlecodeveld.")

add_heading(doc, "3.3 Inloggen met 2FA", 2)
for text in [
    "Bij een volgende login vul je een 6-cijferige code in uit je authenticator-app of uit een ontvangen e-mail.",
    "Als je een herstelcode hebt, kun je die op hetzelfde scherm gebruiken wanneer dat nodig is.",
]:
    add_body_paragraph(doc, text)
add_body_paragraph(doc, "Screenshot toevoegen: het 2FA-controlescherm.")

add_heading(doc, "4. Werkwijze voor chauffeurs", 1)
add_body_paragraph(doc, "Onderstaande werkwijze sluit aan op de manier waarop de applicatie is ingericht.")

for step in [
    "Log in met je e-mailadres, wachtwoord en 2FA-code.",
    "Open de rittenlijst en zoek een open rit of een rit die al aan jou is toegewezen.",
    "Neem telefonisch contact op met de contactpersoon om een afspraak te maken.",
    "Vul de afhaaldatum en afhaaltijd in.",
    "Selecteer jouw naam in het veld 'Chauffeur'.",
    "Gebruik 'Bevestig rit' om een bevestigingsmail te sturen naar de contactpersoon en naar jezelf.",
    "Haal de collecte-opbrengst op en onderteken samen met de contactpersoon het afhaalbewijs.",
    "Stort het geld bij Geldmaat.",
    "Vul na afloop het gestorte bedrag en het aantal gereden kilometers in.",
    "Zet de status van de rit op 'Afgehandeld' en verstuur de afrondingsmail. Voeg daarbij zo nodig een foto van de transactiebon toe.",
    "Open daarna het rapport om je afgeronde ritten en kilometers terug te zien.",
]:
    add_number(doc, step)

add_note_box(
    doc,
    "Praktisch verschil tussen desktop en mobiel",
    [
        "De volledige werkstappen worden uitgevoerd in de uitgebreide rittenlijst.",
        "De mobiele weergave is vooral bedoeld om onderweg snel de ritdetails te bekijken en de contactpersoon te bellen of mailen.",
    ],
)

add_heading(doc, "5. De rittenlijst gebruiken", 1)
add_heading(doc, "5.1 Kleuren in de rittenlijst", 2)
add_bullet(doc, "Wit: de rit is nog niet toegewezen aan een chauffeur.")
add_bullet(doc, "Rood: de rit is wel toegewezen, maar nog niet afgehandeld.")
add_bullet(doc, "Groen: de rit is afgehandeld.")
add_body_paragraph(doc, "Screenshot toevoegen: de rittenlijst met minimaal een witte, rode en groene regel.")

add_heading(doc, "5.2 Welke velden een chauffeur gebruikt", 2)
for text in [
    "Voor chauffeurs zijn vooral de velden voor chauffeur, afhaaldatum, afhaaltijd, gestort bedrag, gereden kilometers en status belangrijk.",
    "De applicatie is zo ingericht dat juist deze velden in de dagelijkse uitvoering nodig zijn.",
]:
    add_body_paragraph(doc, text)

add_heading(doc, "5.3 Rit bevestigen", 2)
for text in [
    "Wanneer de afspraak met de contactpersoon vastligt, gebruik je de knop 'Bevestig rit'.",
    "De applicatie verstuurt dan een bevestigingsmail naar de contactpersoon en een bevestiging naar de chauffeur.",
    "In deze bevestigingen kunnen ook links staan naar een busbriefje en een afhaalbevestiging.",
]:
    add_body_paragraph(doc, text)
add_body_paragraph(doc, "Screenshot toevoegen: de knop 'Bevestig rit' en bij voorkeur het bevestigingsvenster.")

add_heading(doc, "5.4 Rit afronden", 2)
for text in [
    "Na het ophalen en storten vul je het gestorte bedrag en het aantal gereden kilometers in.",
    "Daarna zet je de status op 'Afgehandeld'.",
    "Volgens de werkwijze in de applicatie opent dan het e-mailscherm waarmee de afronding wordt verstuurd. Daar kun je zo nodig een foto van de transactiebon als bijlage toevoegen.",
]:
    add_body_paragraph(doc, text)
add_body_paragraph(doc, "Screenshot toevoegen: het e-mailscherm voor de afronding, inclusief plek voor eventuele bijlagen.")

add_heading(doc, "6. Mobiele weergave", 1)
for text in [
    "Op telefoon komt een chauffeur in de compacte mobiele weergave 'Mijn ritten'.",
    "Hier zie je de ritten die voor jouw account zichtbaar zijn en die nog niet zijn afgehandeld.",
    "Per rit zie je onder andere de contactpersoon, het adres, de geplande datum en tijd, het soort opbrengst, het verwachte bedrag, de chauffeur en de status.",
]:
    add_body_paragraph(doc, text)

add_bullet(doc, "Met 'Open route' open je direct de route in Google Maps.")
add_bullet(doc, "Met 'Bel contactpersoon' bel je direct het opgeslagen telefoonnummer.")
add_bullet(doc, "Met 'Mail contactpersoon' open je direct een e-mail aan de contactpersoon.")
add_bullet(doc, "Met 'Vernieuwen' laad je de rittenlijst opnieuw.")
add_body_paragraph(doc, "Screenshot toevoegen: de mobiele kaartweergave van een rit, inclusief status en knoppen.")

add_heading(doc, "7. Afhaalbewijs en busbriefje", 1)
for text in [
    "Bij het bevestigen van een rit kunnen links worden meegestuurd naar documenten die horen bij de rit.",
    "Een afhaalbevestiging is bedoeld als bewijs van overdracht en bevat ruimte voor handtekeningen van de chauffeur en de contactpersoon.",
    "Een busbriefje kan worden gebruikt als extra document in het proces, afhankelijk van de werkwijze binnen de organisatie.",
]:
    add_body_paragraph(doc, text)
add_body_paragraph(doc, "Screenshot toevoegen: het afhaalbewijs of busbriefje zoals het in de praktijk wordt gebruikt.")

add_heading(doc, "8. Rapport bekijken", 1)
for text in [
    "Via de knop 'Rapport' krijg je een overzicht van ritten met de status 'Afgehandeld'.",
    "Voor chauffeurs toont dit overzicht in elk geval de collectegebieden, de afhaalmomenten, het gestorte bedrag, de gereden kilometers en het declarabele bedrag.",
    "De applicatie berekent het declarabele bedrag op basis van het aantal gereden kilometers.",
    "Volgens de instructie in de applicatie hoeft de chauffeur zelf geen aparte declaratiehandeling te doen; de organisatie handelt dit verder af.",
]:
    add_body_paragraph(doc, text)
add_body_paragraph(doc, "Screenshot toevoegen: het rapport met afgeronde ritten en kilometers.")

add_heading(doc, "9. Voor collega's in de organisatie", 1)
for text in [
    "Collega's met uitgebreide rechten kunnen de applicatie gebruiken om ritten en chauffeurs te beheren.",
    "Zij kunnen ritten toevoegen, ritgegevens aanpassen, chauffeurs toevoegen of verwijderen en overzichten bekijken.",
    "Ook kunnen zij bevestigingsmails versturen en rapporten per chauffeur of voor alle chauffeurs openen.",
]:
    add_body_paragraph(doc, text)

add_bullet(doc, "Nieuwe ritten toevoegen en bestaande ritten wijzigen.")
add_bullet(doc, "Chauffeurs beheren, inclusief naam, postcode, e-mailadres en IBAN.")
add_bullet(doc, "Lat/lon-gegevens opnieuw laten berekenen voor ritten en chauffeurs.")
add_bullet(doc, "Inzicht krijgen in de voortgang: open, toegewezen of afgehandelde ritten.")
add_body_paragraph(doc, "Screenshot toevoegen: het schermdeel voor rittenbeheer en, als relevant, het schermdeel voor chauffeursbeheer.")

add_heading(doc, "10. Veelvoorkomende vragen", 1)
add_heading(doc, "10.1 Ik zie mijn rit niet", 2)
for text in [
    "Controleer of de rit al aan jou is toegewezen of nog open staat.",
    "Gebruik de knop 'Vernieuwen' in de mobiele weergave of laad de rittenlijst opnieuw.",
]:
    add_body_paragraph(doc, text)

add_heading(doc, "10.2 Ik kan op mijn telefoon geen rit afronden", 2)
for text in [
    "De mobiele weergave is vooral bedoeld om ritten te bekijken en contact op te nemen.",
    "Voor de volledige afhandeling is de uitgebreide rittenlijst nodig.",
]:
    add_body_paragraph(doc, text)

add_heading(doc, "10.3 Ik ontvang geen 2FA-code", 2)
for text in [
    "Controleer eerst of je de juiste methode gebruikt: authenticator-app of e-mail.",
    "Controleer bij e-mail ook de map met ongewenste mail.",
    "Lukt het nog steeds niet, neem dan contact op met een collega met beheerrechten.",
]:
    add_body_paragraph(doc, text)

add_heading(doc, "11. Benodigde screenshots", 1)
add_body_paragraph(doc, "Onderstaande lijst kun je gebruiken om later de screenshots op de juiste plaatsen in deze handleiding toe te voegen.")
add_screenshot_table(doc)
add_body_paragraph(doc, "Optioneel: voeg ook een screenshot toe van het chauffeursbeheer als je het organisatorische beheeronderdeel visueel wilt toelichten.")

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
